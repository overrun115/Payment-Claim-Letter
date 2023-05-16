import pandas as pd
from docxtpl import DocxTemplate
import sys
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from num2words import num2words

def ConvertToWords(Number):
    Integer, Decimal = str(round(Number, 2)).split(".")
    Integer = int(Integer)
    Decimal = Decimal.ljust(2, '0')
    IntegerLetters = num2words(Integer, lang='en')   
    if Decimal == "00":
        DecimalLetters = "zero"
    else:
        DecimalLetters = f"{Decimal}/100"
    return f"{IntegerLetters} with {DecimalLetters}"

if getattr(sys, 'frozen', False):
    CurrentPath = os.path.dirname(sys.executable)
else:
    CurrentPath = os.path.dirname(os.path.abspath(__file__))

FolderPath = CurrentPath
DocPath = os.path.join(FolderPath, 'PaymentClaimLetter.docx')
DebtPath = os.path.join(FolderPath, 'DebtDetail.csv')

doc = DocxTemplate(DocPath)

DfDebt = pd.read_csv(DebtPath, sep=";", decimal='.', encoding='utf-8')

Df = DfDebt.loc[:, ['Account', 'Name']].drop_duplicates()

DfBusiness = DfDebt.groupby(['Account', 'Unit'])['Amount'].sum().reset_index()
idx = DfBusiness.groupby('Account')['Amount'].idxmax()
DfBusiness = DfBusiness.loc[idx, ['Account', 'Unit']].set_index('Account')
Df['Business'] = DfBusiness['Unit']

def GetBusiness(row):
    Account = row['Account']
    if Account in DfBusiness.index:
        return DfBusiness.loc[Account, 'Unit']
    else:
        return ''

Df['Business'] = Df.apply(GetBusiness, axis=1)
Df['Debt'] = Df.apply(lambda row: DfDebt.loc[DfDebt['Account'] == row['Account'], 'Amount'].sum(), axis=1)

for index, row in Df.iterrows():
    row_dict = row.to_dict()
    TableData = DfDebt[DfDebt['Account'] == row['Account']].to_dict(orient='records')

    context = {'Name': row['Name'],
            'Business': row['Business'],
            'Debt': "$"+"{:,.2f}".format(float(row['Debt'])), 
            'Account': row['Account'],
            'DebtLetter': ConvertToWords(float(row['Debt']))}

    Table = []
    Header = ['Document', 'Unit', 'Class', 'Doc. date', 'Due date', 'Amount']
    Table.append(Header)
    for item in TableData:
        row = []
        row.append(item['Document'])
        row.append(item['Unit'])
        row.append(item['Class'])
        row.append(item['Doc. date'].replace('.', '/'))
        row.append(item['Due date'].replace('.', '/'))
        row.append("{:,.2f}".format(float(item['Amount']))) 
        Table.append(row)

    doc.render(context)

    for i, row in enumerate(Table):
        for j, cell in enumerate(row):

            while i >= len(doc.tables[0].rows):
                doc.tables[0].add_row()
            while j >= len(doc.tables[0].row_cells(i)):
                doc.tables[0].rows[i].add_cell()
            doc.tables[0].cell(i, j).text = str(cell)
        
            TableCell = doc.tables[0].cell(i, j)
            TableCell.text = str(cell)
            if j == len(Header) - 1: 
                TableCell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.tables[0].autofit = True
    
    FilePath = os.path.join(FolderPath, f"{row_dict['Account']}-{row_dict['Name'].upper()}.docx")
    doc.save(FilePath)